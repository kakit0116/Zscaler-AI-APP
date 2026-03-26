import os
import glob
from dotenv import load_dotenv

load_dotenv()

#from langchain_openai import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredMarkdownLoader

def load_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs if p.text.strip()])

def load_file(path: str):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return PyPDFLoader(path).load()
    if ext in [".md", ".markdown"]:
        return UnstructuredMarkdownLoader(path).load()
    if ext in [".txt", ".log"]:
        return TextLoader(path, encoding="utf-8").load()
    if ext == ".docx":
        from langchain_core.documents import Document
        return [Document(page_content=load_docx(path), metadata={"source": path})]
    return []

def collect_documents(data_dir: str):
    patterns = ["**/*.pdf", "**/*.md", "**/*.markdown", "**/*.txt", "**/*.log", "**/*.docx"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(data_dir, p), recursive=True))

    all_docs = []
    for f in files:
        docs = load_file(f)
        for d in docs:
            d.metadata = d.metadata or {}
            d.metadata.setdefault("source", f)
        all_docs.extend(docs)
    return all_docs

def main():
    data_dir = os.getenv("DATA_DIR", "data")
    out_dir = os.getenv("FAISS_DIR", "faiss_store")

    docs = collect_documents(data_dir)
    if not docs:
        raise RuntimeError(f"No docs found in ./{data_dir}. Put files under ./data")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=120,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    # Use the same embedding model as app.py for consistency
    embeddings = HuggingFaceEmbeddings(
        model_name=os.getenv("HF_EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
    )

    db = FAISS.from_documents(chunks, embeddings)
    db.save_local(out_dir)

    print(f"✅ Loaded {len(docs)} docs → {len(chunks)} chunks")
    print(f"✅ Saved FAISS index to: {out_dir}/")

if __name__ == "__main__":
    main()

